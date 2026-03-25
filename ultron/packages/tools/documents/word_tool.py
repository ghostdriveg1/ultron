import os
import tempfile
from docx import Document

from packages.tools.base_tool import BaseTool
from packages.tools.schemas.word_schema import WordInput, WordOutput
from .fast_io_client import FastIOClient

class WordTool(BaseTool):
    """Tool to create Word documents (.docx)."""
    input_schema = WordInput
    output_schema = WordOutput

    def __init__(self):
        super().__init__(
            name="create_word_doc",
            description="Creates a Word document from text content.",
            permission_level="ALWAYS_ALLOWED"
        )
        self.fast_io = FastIOClient()

    async def execute(self, params: WordInput) -> WordOutput:
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, params.output_filename)
            
            # Create Word Doc
            doc = Document()
            doc.add_heading(params.title, 0)
            
            for line in params.content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
                    
            doc.save(file_path)
            
            # Upload
            fast_io_url = await self.fast_io.upload(file_path, params.output_filename)
            
            # Mock GitHub URL
            github_url = f"https://github.com/mock/outputs/{params.output_filename}"
            
            return WordOutput(
                fast_io_url=fast_io_url,
                github_url=github_url
            )
